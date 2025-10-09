# tests/test_ingest.py
# -*- coding: utf-8 -*-
import sys
from pathlib import Path

# 让 "scripts" 可导入
ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.ingest import clean_text, detect_and_read  # noqa: E402


def test_clean_text_normalizes_newlines_and_compresses_blank_lines():
    raw = "行1\r\n\r\n\r\n行2\r行3  \n\n\n"
    out = clean_text(raw)
    # 统一换行为 \n；多空行压缩到 1 行；去掉行尾空格
    assert out == "行1\n\n行2\n行3"


def test_detect_and_read_handles_utf8_bom(tmp_path: Path):
    p = tmp_path / "bom.txt"
    content = "第一章 UTF8-BOM\n\n这是内容A。\n"
    p.write_text(content, encoding="utf-8-sig")  # 写入带BOM
    text, enc, size = detect_and_read(p)
    assert "第一章 UTF8-BOM" in text
    assert len(text) == len(content)
    assert "utf-8" in enc.replace("_", "-").lower()
    assert size > 0


def test_detect_and_read_handles_gbk(tmp_path: Path):
    p = tmp_path / "gbk.txt"
    content = "第一章 GBK\n\n这是内容B（逗号，句号。中文标点测试）。"
    p.write_text(content, encoding="gbk", errors="strict")
    text, enc, size = detect_and_read(p)
    assert "第一章 GBK" in text
    assert "内容B" in text
    assert size > 0
    assert any(k in enc.lower() for k in ["gbk", "cp936", "utf-8"])


def test_detect_and_read_docx(tmp_path: Path):
    from docx import Document
    p = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("段落一")
    doc.add_paragraph("")  # 空行
    doc.add_paragraph("段落二——带中文标点，测试")
    doc.save(p)

    text, enc, size = detect_and_read(p)
    assert "段落一" in text and "段落二" in text
    assert enc == "docx"
    assert size > 0
